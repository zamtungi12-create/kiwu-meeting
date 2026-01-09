import streamlit as st
import pandas as pd
import gspread
from datetime import datetime
from io import BytesIO
import os # 파일 존재 여부 확인용

# [필수] 워드 파일 생성을 위한 라이브러리
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- [0] 부서 순서 정의 (고정 리스트) ---
DEPT_ORDER = [
    "교목실", "감사팀", "기획팀", "미래전략센터", "혁신지원사업단", 
    "교무수업팀", "교무인사팀", "교육혁신센터", "학사학위센터", 
    "학생복지팀", "장애학생지원센터", "학생상담센터", "사회공헌센터", 
    "커뮤니케이션팀", "입학지원팀", "취창업진로지원센터", "산학운영팀", 
    "RISE사업단", "현장실습지원센터", "일학습병행공동훈련센터", 
    "총무팀", "시설안전팀", "국제교육팀", "글로벌커리어센터", 
    "글로벌인재정주지원센터", "평생교육원", "도서관", "전산정보원", "SG캠퍼스사업단"
]

# --- [1] 기본 설정 및 디자인 ---
st.set_page_config(page_title="KIWU Smart Meeting", page_icon="🎓", layout="wide")

st.markdown("""
    <style>
    /* 전체 폰트 적용 */
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700;900&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Noto Sans KR', sans-serif;
    }
    
    .stApp { background-color: #f8f9fa; }

    /* 헤더 스타일 */
    .main-header { 
        font-size: 2.2rem; color: #003478; font-weight: 900; 
        margin-top: 10px; margin-bottom: 5px; 
    }
    .sub-header {
        font-size: 1.0rem; color: #666; margin-bottom: 25px;
    }

    /* 카드 박스 */
    .card-box { 
        background-color: white; padding: 20px 10px; border-radius: 10px; 
        border: 1px solid #edf2f7; 
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1); 
        text-align: center; border-top: 4px solid #003478; 
    }
    .card-box h5 { margin: 0; font-size: 0.9rem; color: #718096; }
    .card-box h2 { margin: 5px 0 0 0; font-size: 1.8rem; font-weight: 700; color: #2d3748; }
    
    .admin-box { 
        background-color: #ebf8ff; padding: 20px; border-radius: 10px; border: 1px solid #bee3f8; 
    }

    /* [중요] HTML 테이블 스타일 정의 */
    .kiwu-table-container {
        overflow-x: auto;
    }
    table.kiwu-table {
        width: 100%;
        border-collapse: collapse;
        margin: 10px 0;
        background-color: white;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    /* 헤더 스타일 */
    table.kiwu-table th {
        background-color: #f0f2f6;
        color: #003478;
        font-size: 15px;       
        font-weight: 900;      
        text-align: center;    
        padding: 15px 10px;
        border-bottom: 3px solid #003478; 
        white-space: nowrap;
    }
    /* 데이터 셀 스타일 */
    table.kiwu-table td {
        padding: 12px 10px;
        border-bottom: 1px solid #e2e8f0;
        text-align: center;    
        font-size: 15px;
        color: #333;
    }
    /* 업무내용 컬럼(3번째)만 좌측 정렬 */
    table.kiwu-table td:nth-child(3) {
        text-align: left;
        min-width: 300px;
    }

    @media print {
        .stSidebar, header, footer, .no-print { display: none !important; }
        .print-only { display: block !important; }
        .stApp { background-color: white !important; }
    }
    </style>
""", unsafe_allow_html=True)

# --- [2] 구글 시트 연결 함수 ---
@st.cache_resource
def get_connection():
    try:
        if "gcp_service_account" in st.secrets:
            creds_dict = st.secrets["gcp_service_account"]
            gc = gspread.service_account_from_dict(creds_dict)
        else:
            gc = gspread.service_account(filename='service_account.json')
    except Exception:
        gc = gspread.service_account(filename='service_account.json')
    return gc

def get_google_sheet(sheet_name):
    gc = get_connection()
    doc = gc.open("경인여대 스마트회의 DB")
    return doc.worksheet(sheet_name)

# --- [3] 스타일링된 HTML 테이블 생성 함수 ---
def render_styled_table(df):
    html = df.to_html(index=False, classes='kiwu-table', escape=False)
    st.markdown(f'<div class="kiwu-table-container">{html}</div>', unsafe_allow_html=True)

# --- [4] 워드 파일 생성 함수 ---
def create_docx(df, title_text):
    doc = Document()
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("-" * 50)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["부서", "구분", "내용", "상태", "기한", "담당자"]
    
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(12) 

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['부서명'])
        row_cells[1].text = str(row['구분'])
        row_cells[2].text = str(row['업무내용'])
        row_cells[3].text = str(row['진행상태'])
        row_cells[4].text = str(row['마감기한'])
        row_cells[5].text = str(row['담당자'])
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- [5] 전체화면 팝업 함수 ---
@st.dialog("🔍 전체 안건 확대 보기", width="large")
def show_fullscreen_table(df):
    st.markdown("### 📋 전체 안건 목록")
    render_styled_table(df) 
    if st.button("닫기"):
        st.rerun()

# --- [6] 사이드바 메뉴 (로고 적용 부분) ---
with st.sidebar:
    # [수정] 로고 이미지 표시 로직 (파일이 있으면 이미지, 없으면 텍스트)
    # logo.png 파일이 프로젝트 폴더에 있어야 합니다.
    if os.path.exists("logo.png"):
        st.image("logo.png", use_container_width=True)
    elif os.path.exists("logo.jpg"):
        st.image("logo.jpg", use_container_width=True)
    else:
        # 이미지가 없으면 텍스트로 깔끔하게 표시
        st.markdown("## 🎓 KIWU Admin")
    
    menu = st.radio("메뉴 선택", [
        "📊 금주 현황 (Current)", 
        "📝 안건 등록 (Input)", 
        "🛠️ 수정/삭제 (Edit)", 
        "🗄️ 지난 기록 (History)", 
        "🖨️ 회의록 다운로드 (Export)", 
        "⚙️ 관리자 (Admin)"
    ])
    st.markdown("---")
    if st.button("🔄 새로고침"):
        st.rerun()

# --- [7] 기능: 금주 현황 ---
if menu == "📊 금주 현황 (Current)":
    current_hour = datetime.now().hour 
    caption_text = "경인여자대학교의 힘찬 하루 ☀️" if 6 <= current_hour < 18 else "경인여자대학교의 빛나는 열정 🌙"
    st.caption(caption_text)

    st.markdown('<div class="main-header">🎓 경인여자대학교 전략회의</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="sub-header">📅 기준일: {datetime.now().strftime("%Y년 %m월 %d일")} | 종이 없는 스마트 회의 시스템</div>', unsafe_allow_html=True)
    
    try:
        sheet = get_google_sheet("Current")
        data = sheet.get_all_records()
        df = pd.DataFrame(data)

        submitted_depts = []
        if not df.empty:
            submitted_depts = df['부서명'].unique()
        unsubmitted_list = [d for d in DEPT_ORDER if d not in submitted_depts]

        if unsubmitted_list:
            with st.expander(f"🚨 미제출 부서 현황: 총 {len(unsubmitted_list)}개 부서 (클릭하여 명단 확인)", expanded=False):
                st.markdown(f"""
                <div style='background-color: #fff5f5; padding: 15px; border-radius: 5px; border-left: 5px solid #fc8181;'>
                    <p style='color: #c53030; font-weight: bold; margin: 0;'>{', '.join(unsubmitted_list)}</p>
                    <p style='font-size: 0.85rem; color: #718096; margin-top: 8px; margin-bottom: 0;'>
                        ※ 원활한 회의 진행을 위해 빠른 입력을 부탁드립니다.
                    </p>
                </div>
                """, unsafe_allow_html=True)
        else:
            if not df.empty:
                st.balloons()
                st.success("🎉 모든 부서가 안건 제출을 완료했습니다!")

        if not df.empty:
            cnt_total = len(df)
            cnt_ing = len(df[df['진행상태'] == '진행중'])
            cnt_plan = len(df[df['진행상태'] == '예정'])
            cnt_done = len(df[df['진행상태'] == '완료'])
            cnt_delay = len(df[df['진행상태'] == '지연'])

            c1, c2, c3, c4, c5 = st.columns(5)
            with c1: st.markdown(f'<div class="card-box"><h5>전체 안건</h5><h2>{cnt_total}</h2></div>', unsafe_allow_html=True)
            with c2: st.markdown(f'<div class="card-box" style="border-top-color: #e67e22;"><h5>진행중</h5><h2 style="color:#e67e22;">{cnt_ing}</h2></div>', unsafe_allow_html=True)
            with c3: st.markdown(f'<div class="card-box" style="border-top-color: #3182ce;"><h5>예정</h5><h2 style="color:#3182ce;">{cnt_plan}</h2></div>', unsafe_allow_html=True)
            with c4: st.markdown(f'<div class="card-box" style="border-top-color: #38a169;"><h5>완료</h5><h2 style="color:#38a169;">{cnt_done}</h2></div>', unsafe_allow_html=True)
            with c5: st.markdown(f'<div class="card-box" style="border-top-color: #e53e3e;"><h5>지연</h5><h2 style="color:#e53e3e;">{cnt_delay}</h2></div>', unsafe_allow_html=True)
            
            st.markdown("---")
            
            col_filter, col_btn = st.columns([0.85, 0.15]) 
            
            with col_filter:
                with st.expander("🔍 부서별 필터링 옵션 (클릭하여 펼치기)", expanded=False):
                    unique_depts = df['부서명'].unique()
                    sorted_depts = [d for d in DEPT_ORDER if d in unique_depts]
                    others = [d for d in unique_depts if d not in DEPT_ORDER]
                    final_dept_list = sorted_depts + others
                    selected_dept = st.multiselect("보고 싶은 부서를 선택하세요:", final_dept_list, default=final_dept_list)
            
            if selected_dept:
                filtered_df = df[df['부서명'].isin(selected_dept)]
                filtered_df['부서명'] = pd.Categorical(filtered_df['부서명'], categories=DEPT_ORDER + others, ordered=True)
                filtered_df = filtered_df.sort_values('부서명')
                
                display_df = filtered_df.drop(columns=['비밀번호']) if '비밀번호' in filtered_df.columns else filtered_df

                render_styled_table(display_df)

                with col_btn:
                    st.write("") 
                    if st.button("🖥️ 크게 보기", type="secondary"):
                        show_fullscreen_table(display_df)
            else:
                st.info("선택된 부서가 없습니다. 필터를 확인해주세요.")
        else:
            st.info("👋 아직 등록된 안건이 없습니다.")

    except Exception as e:
        st.error(f"오류: {e}")

# --- [8] 기능: 안건 등록 ---
elif menu == "📝 안건 등록 (Input)":
    st.markdown('<div class="main-header">📝 안건 등록</div>', unsafe_allow_html=True)
    st.info("입력된 내용은 '이번 주 현황'에 즉시 반영됩니다.")

    with st.form("input_form", clear_on_submit=True):
        col_a, col_b = st.columns(2)
        with col_a:
            input_dept = st.selectbox("부서", DEPT_ORDER)
            input_type = st.selectbox("구분", ["주요현안", "일반보고", "협조요청"])
        with col_b:
            input_status = st.selectbox("상태", ["진행중", "완료", "지연", "예정"])
            input_date = st.date_input("마감 기한")
        
        input_content = st.text_area("업무 내용", height=100)
        col_c, col_d = st.columns(2)
        with col_c: input_name = st.text_input("담당자")
        with col_d: input_note = st.text_input("비고")
        
        st.markdown("---")
        input_pw = st.text_input("비밀번호(수정/삭제용 숫자 4자리)", type="password", max_chars=4)
        
        if st.form_submit_button("💾 등록하기", type="primary"):
            if not input_pw:
                st.warning("비밀번호를 입력해주세요!")
            else:
                try:
                    sheet = get_google_sheet("Current")
                    now = datetime.now().strftime("%Y-%m-%d %H:%M")
                    sheet.append_row([now, input_dept, input_type, input_content, input_status, str(input_date), input_name, input_note, input_pw])
                    st.success("등록되었습니다!")
                except Exception as e:
                    st.error(f"저장 실패: {e}")

# --- [9] 기능: 수정/삭제 ---
elif menu == "🛠️ 수정/삭제 (Edit)":
    st.markdown('<div class="main-header">🛠️ 안건 수정 및 삭제</div>', unsafe_allow_html=True)
    try:
        sheet = get_google_sheet("Current")
        data = sheet.get_all_records()
        df = pd.DataFrame(data)
        if df.empty:
            st.info("수정할 데이터가 없습니다.")
        else:
            dept_list_for_edit = sorted(df['부서명'].unique())
            edit_dept = st.selectbox("부서를 선택하세요", dept_list_for_edit)
            target_df = df[df['부서명'] == edit_dept]
            
            if not target_df.empty:
                task_options = target_df.apply(lambda x: f"[{x['입력일시']}] {str(x['업무내용'])[:20]}...", axis=1)
                selected_task_idx = st.selectbox("안건을 선택하세요", task_options.index, format_func=lambda x: task_options[x])
                selected_row = df.loc[selected_task_idx]
                
                st.info(f"선택: {selected_row['업무내용']}")
                chk_pw = st.text_input("비밀번호 확인", type="password")
                
                if st.button("확인"):
                    if str(selected_row.get('비밀번호', '')) == str(chk_pw):
                        st.session_state['auth_success'] = True
                        st.session_state['target_idx'] = selected_task_idx 
                        st.success("인증 성공")
                    else:
                        st.error("비밀번호 불일치")
                        st.session_state['auth_success'] = False
                
                if st.session_state.get('auth_success', False) and st.session_state.get('target_idx') == selected_task_idx:
                    st.divider()
                    st.subheader("내용 수정")
                    with st.form("edit_form"):
                        def safe_index(lst, val): return lst.index(val) if val in lst else 0
                        e_type = st.selectbox("구분", ["주요현안", "일반보고", "협조요청"], index=safe_index(["주요현안", "일반보고", "협조요청"], selected_row['구분']))
                        e_status = st.selectbox("상태", ["진행중", "완료", "지연", "예정"], index=safe_index(["진행중", "완료", "지연", "예정"], selected_row['진행상태']))
                        e_content = st.text_area("업무 내용", value=selected_row['업무내용'])
                        e_note = st.text_input("비고", value=selected_row['비고'])
                        
                        c1, c2 = st.columns(2)
                        with c1: update_btn = st.form_submit_button("수정 저장", type="primary")
                        with c2: delete_btn = st.form_submit_button("🗑️ 삭제하기")
                        
                        real_row_num = selected_task_idx + 2 
                        if update_btn:
                            sheet.update_cell(real_row_num, 3, e_type)
                            sheet.update_cell(real_row_num, 4, e_content)
                            sheet.update_cell(real_row_num, 5, e_status)
                            sheet.update_cell(real_row_num, 8, e_note)
                            st.success("수정 완료! 새로고침 해주세요.")
                            del st.session_state['auth_success']
                        if delete_btn:
                            sheet.delete_rows(real_row_num)
                            st.success("삭제 완료! 새로고침 해주세요.")
                            del st.session_state['auth_success']
            else:
                st.warning("해당 부서에 등록된 안건이 없습니다.")
    except Exception as e:
        st.error(f"오류: {e}")

# --- [10] 기능: 지난 기록 ---
elif menu == "🗄️ 지난 기록 (History)":
    st.markdown('<div class="main-header">🗄️ 지난 회의 기록</div>', unsafe_allow_html=True)
    try:
        sheet = get_google_sheet("History")
        data = sheet.get_all_records()
        df = pd.DataFrame(data)
        if not df.empty:
            meeting_dates = list(df['회차정보'].unique())
            selected_date = st.selectbox("회차 선택:", meeting_dates)
            
            history_df = df[df['회차정보'] == selected_date]
            unique_depts_hist = df['부서명'].unique()
            others_hist = [d for d in unique_depts_hist if d not in DEPT_ORDER]
            history_df['부서명'] = pd.Categorical(history_df['부서명'], categories=DEPT_ORDER + others_hist, ordered=True)
            history_df = history_df.sort_values('부서명')
            
            render_styled_table(history_df)
        else:
            st.warning("보관된 기록이 없습니다.")
    except Exception as e:
        st.error(f"오류: {e}")

# --- [11] 기능: 회의록 다운로드 ---
elif menu == "🖨️ 회의록 다운로드 (Export)":
    st.markdown('<div class="main-header">🖨️ 회의록 생성 및 다운로드</div>', unsafe_allow_html=True)
    
    export_target = st.radio("출력할 대상을 선택하세요:", ["금주 안건 (Current)", "지난 기록 (History)"], horizontal=True)
    target_df = pd.DataFrame()
    report_title = ""

    try:
        if "금주" in export_target:
            sheet = get_google_sheet("Current")
            data = sheet.get_all_records()
            target_df = pd.DataFrame(data)
            report_title = f"{datetime.now().strftime('%Y-%m-%d')} 전략회의 안건"
        else:
            sheet = get_google_sheet("History")
            data = sheet.get_all_records()
            all_hist_df = pd.DataFrame(data)
            if not all_hist_df.empty:
                meeting_dates = list(all_hist_df['회차정보'].unique())
                selected_date = st.selectbox("출력할 회차를 선택하세요:", meeting_dates)
                target_df = all_hist_df[all_hist_df['회차정보'] == selected_date]
                report_title = f"{selected_date} 회의록"
        
        if not target_df.empty:
            unique_depts = target_df['부서명'].unique()
            sorted_depts = [d for d in DEPT_ORDER if d in unique_depts]
            others = [d for d in unique_depts if d not in DEPT_ORDER]
            target_df['부서명'] = pd.Categorical(target_df['부서명'], categories=DEPT_ORDER + others, ordered=True)
            target_df = target_df.sort_values('부서명')

            cols_to_show = ['부서명', '구분', '업무내용', '진행상태', '마감기한', '담당자']
            final_df = target_df[cols_to_show] if set(cols_to_show).issubset(target_df.columns) else target_df

            st.divider()
            st.subheader(f"📄 미리보기: {report_title}")
            render_styled_table(final_df)

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("### 📥 Word 다운로드")
                docx_file = create_docx(final_df, report_title)
                st.download_button(
                    label="워드 파일 다운로드 (.docx)",
                    data=docx_file,
                    file_name=f"{report_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with c2:
                st.markdown("### 🖨️ 인쇄 / PDF 저장")
                html_table = final_df.to_html(index=False, classes='kiwu-table', escape=False)
                html_content = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: 'Malgun Gothic', sans-serif; padding: 20px; }}
                        h1 {{ text-align: center; color: #003478; }}
                        .date {{ text-align: right; color: #666; margin-bottom: 20px; }}
                        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12px; }}
                        th, td {{ border: 1px solid #444; padding: 8px; }}
                        th {{ 
                            background-color: #f2f2f2; 
                            text-align: center !important; 
                            font-weight: 900 !important; 
                            font-size: 18px !important; 
                            color: #003478;
                            padding: 10px;
                        }}
                        td {{ text-align: center; font-size: 14px; }}
                        td:nth-child(3) {{ text-align: left; }}
                    </style>
                </head>
                <body>
                    <h1>{report_title}</h1>
                    <div class="date">출력일: {datetime.now().strftime('%Y-%m-%d')}</div>
                    {html_table}
                </body>
                </html>
                """
                with st.expander("👁️ 인쇄용 뷰 열기 (클릭)"):
                    st.components.v1.html(html_content, height=600, scrolling=True)
                    st.info("💡 위 표 위에서 마우스 오른쪽 버튼 -> '프레임 인쇄' 또는 이 화면 전체를 'Ctrl+P'로 인쇄하세요.")
    except Exception as e:
        st.error(f"데이터 처리 중 오류 발생: {e}")

# --- [12] 기능: 관리자 ---
elif menu == "⚙️ 관리자 (Admin)":
    st.markdown('<div class="main-header">⚙️ 관리자 페이지</div>', unsafe_allow_html=True)
    password = st.text_input("관리자 비밀번호", type="password")

    try:
        real_password = st.secrets["admin"]["password"] if "admin" in st.secrets else "1234"
    except:
        real_password = "1234"
    
    if password == real_password:
        st.success("✅ 접속 완료")
        st.markdown("""
        <div class="admin-box">
            <h4>🔴 주간 회의 마감 (Data Closing)</h4>
            <p><b>[Current]</b> 데이터를 <b>[History]</b>로 이동하고 초기화합니다.</p>
        </div>
        """, unsafe_allow_html=True)
        
        meeting_name = st.text_input("마감할 회차 이름 (예: 2026-01-08 정기회의)")
        confirm_close = st.checkbox("⚠️ 데이터 마감 확인")
        
        if st.button("🚀 마감 실행"):
            if not confirm_close or not meeting_name:
                st.warning("회차 이름과 확인 체크박스를 모두 입력해주세요.")
            else:
                try:
                    cur_sheet = get_google_sheet("Current")
                    his_sheet = get_google_sheet("History")
                    data = cur_sheet.get_all_values()
                    
                    if len(data) <= 1:
                        st.warning("데이터 없음")
                    else:
                        records = data[1:]
                        history_records = []
                        for row in records:
                            safe_row = row[:-1]
                            safe_row.insert(0, meeting_name)
                            history_records.append(safe_row)
                        his_sheet.append_rows(history_records)
                        cur_sheet.batch_clear(["A2:Z1000"])
                        st.balloons()
                        st.success("✅ 마감 완료")
                except Exception as e:
                    st.error(f"오류: {e}")
    elif password:
        st.error("비밀번호 불일치")